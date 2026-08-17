# -*- coding: utf-8 -*-
"""
docx_builder.py
Reading Vocabulary Maker MASTER MANUAL v1.0 - 9장 "Word 문서 디자인 원칙" 구현.

디자인 잠금 항목 (AI가 임의로 바꾸지 않고 코드가 고정):
  - 세로(portrait) 페이지 방향
  - 글씨체 및 글자 크기 (맑은 고딕)
  - 페이지 여백
  - 제목 위치 (상단 중앙)
  - footer ("Hard work pays off.")
  - 표의 기본 구조와 너비, 실선 셀 구분
  - No. 칸은 좁게 유지
  - 시험지는 학생 답안 작성 공간 확보
  - 정답·해설지는 번호/정답 열은 좁게, 해설 영역은 넓게
"""

import io

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "맑은 고딕"
FOOTER_TEXT = "Hard work pays off."
CHOICE_LABELS = ["①", "②", "③", "④"]


# ---------------------------------------------------------------------------
# 공통 헬퍼
# ---------------------------------------------------------------------------

def _set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _new_document():
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = FONT_NAME
    normal_style.font.size = Pt(10.5)
    rpr = normal_style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT_NAME)

    footer = section.footer
    footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(FOOTER_TEXT)
    _set_font(footer_run, size=9)

    return doc


def _add_title(doc, title_text, subtitle_text=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_text)
    _set_font(run, size=18, bold=True)

    if subtitle_text:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle_text)
        _set_font(run2, size=11)

    doc.add_paragraph()


def _set_cell_border(cell, spec):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(spec.get("sz", 6)))
        el.set(qn("w:color"), spec.get("color", "999999"))


def _apply_table_borders(table):
    spec = {"sz": 6, "val": "single", "color": "999999"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, spec)


def _shade_cell(cell, color_hex="E7E7E7"):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcpr.append(shd)


def _set_cell_text(cell, text, size=9.5, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run("" if text is None else str(text))
    _set_font(run, size=size, bold=bold)


def _set_column_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)
    for idx, width in enumerate(widths_cm):
        if idx < len(table.columns):
            table.columns[idx].width = Cm(width)


# ---------------------------------------------------------------------------
# 1) Vocabulary List
# ---------------------------------------------------------------------------

def build_vocab_list_docx(title, vocab_list):
    doc = _new_document()
    _add_title(doc, title, "Vocabulary List")

    headers = ["No.", "Word / Expression", "Pronunciation", "POS", "Meaning",
               "Synonym", "Antonym", "New Example", "Source", "Imp."]
    widths = [0.9, 2.3, 1.9, 1.1, 2.6, 1.4, 1.4, 3.0, 1.4, 0.9]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr_cells[i], h, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(hdr_cells[i])

    for i, item in enumerate(vocab_list, start=1):
        row_cells = table.add_row().cells
        values = [
            i,
            item.get("word", ""),
            item.get("pronunciation", ""),
            item.get("pos", ""),
            item.get("meaning", ""),
            item.get("synonym", "-") or "-",
            item.get("antonym", "-") or "-",
            item.get("new_example", ""),
            item.get("source", ""),
            item.get("importance", ""),
        ]
        for col_idx, v in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 3, 9) else None
            _set_cell_text(row_cells[col_idx], v, align=align)

    _set_column_widths(table, widths)
    _apply_table_borders(table)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ---------------------------------------------------------------------------
# 2) Vocabulary Review Test
# ---------------------------------------------------------------------------

SECTION_DEFS = [
    ("Part 1. Reading Review (No. 1-12)", "reading_review"),
    ("Part 2. Vocabulary Transfer (No. 13-21)", "vocabulary_transfer"),
    ("Part 3. English Definition (No. 22-26)", "english_definition"),
    ("Part 4. Vocabulary Relations (No. 27-30)", "vocabulary_relations"),
]


def build_test_docx(title, questions, difficulty):
    doc = _new_document()
    _add_title(doc, title, f"Vocabulary Review Test  (Level: {difficulty})")

    info_p = doc.add_paragraph()
    info_run = info_p.add_run("Name: ______________________        Class: __________        Score:  _______ / 30")
    _set_font(info_run, size=10.5)
    doc.add_paragraph()

    qmap = {q["no"]: q for q in questions if q.get("no") is not None}

    for heading, qtype in SECTION_DEFS:
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(heading)
        _set_font(h_run, size=13, bold=True)

        nums = sorted(no for no, q in qmap.items() if q.get("type") == qtype)
        for no in nums:
            q = qmap[no]
            q_p = doc.add_paragraph()
            q_run = q_p.add_run(f"{no}.  {q.get('question_text', '')}")
            _set_font(q_run, size=11)

            if qtype in ("english_definition", "vocabulary_relations"):
                choices = q.get("choices") or []
                c_p = doc.add_paragraph()
                c_p.paragraph_format.left_indent = Cm(0.6)
                choice_text = "     ".join(
                    f"{CHOICE_LABELS[i]} {c}" for i, c in enumerate(choices) if i < len(CHOICE_LABELS)
                )
                c_run = c_p.add_run(choice_text)
                _set_font(c_run, size=10.5)
            else:
                a_p = doc.add_paragraph()
                a_p.paragraph_format.left_indent = Cm(0.6)
                a_run = a_p.add_run("Answer:  " + "_" * 28)
                _set_font(a_run, size=10.5)

            doc.add_paragraph()

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ---------------------------------------------------------------------------
# 3) Answer & Explanation
# ---------------------------------------------------------------------------

def build_answer_docx(title, questions, vocab_list):
    doc = _new_document()
    _add_title(doc, title, "Answer & Explanation")

    headers = ["No.", "Answer", "Meaning", "Explanation"]
    widths = [1.2, 2.8, 2.8, 10.2]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr_cells[i], h, size=9.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(hdr_cells[i])

    meaning_lookup = {
        str(v.get("word", "")).strip().lower(): v.get("meaning", "") for v in vocab_list
    }

    for q in sorted((q for q in questions if q.get("no") is not None), key=lambda x: x["no"]):
        row_cells = table.add_row().cells

        if q.get("type") in ("english_definition", "vocabulary_relations"):
            choices = q.get("choices") or []
            ai = q.get("answer_index")
            if isinstance(ai, int) and 0 <= ai < len(choices):
                answer_display = f"{CHOICE_LABELS[ai]} {choices[ai]}"
            else:
                answer_display = q.get("answer", "")
        else:
            answer_display = q.get("answer", "")

        target = str(q.get("target_word", "")).strip().lower()
        meaning = meaning_lookup.get(target, "")

        _set_cell_text(row_cells[0], q["no"], align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row_cells[1], answer_display)
        _set_cell_text(row_cells[2], meaning)
        _set_cell_text(row_cells[3], q.get("explanation", ""))

    _set_column_widths(table, widths)
    _apply_table_borders(table)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# -*- coding: utf-8 -*-
"""
Reading Vocabulary Maker
Streamlit Community Cloud 전용 앱

기능
- 독해 PDF/이미지 업로드
- 페이지 범위 선택
- 지문 난이도/어휘 밀도에 따라 어휘 수 자동 조절
- 단어장 / 4유형 시험지 / 정답·해설지 생성
- 문제별 쉽게/어렵게 재생성
- Word(.docx) 다운로드

중요
- Flask를 사용하지 않습니다.
- app.run(), port, gunicorn이 없습니다.
- OpenAI API 키는 Streamlit Secrets의 OPENAI_API_KEY에서 읽습니다.
"""

from __future__ import annotations

import base64
import io
import json
import os
import random
import re
from copy import deepcopy
from typing import Any, Dict, List, Tuple

import fitz  # PyMuPDF
import streamlit as st
from openai import OpenAI
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

# ---------------------------------------------------------------------
# Streamlit / OpenAI
# ---------------------------------------------------------------------

st.set_page_config(
    page_title="Reading Vocabulary Maker",
    page_icon="📘",
    layout="wide",
)

try:
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
except Exception:
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

try:
    DEFAULT_MODEL = st.secrets.get("OPENAI_MODEL", "gpt-4.1-mini")
except Exception:
    DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")

if not OPENAI_API_KEY:
    st.error(
        "OPENAI_API_KEY가 없습니다. "
        "Streamlit → Manage app → Settings → Secrets에 "
        'OPENAI_API_KEY = "..." 형식으로 등록해 주세요.'
    )
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)


# ---------------------------------------------------------------------
# 출제 매뉴얼
# ---------------------------------------------------------------------

SYSTEM_MANUAL = r"""
당신은 한국의 영어 독해 수업을 위한 'Reading Vocabulary Maker'이다.
목표는 단순 암기형 단어 시험이 아니라, 학생이 실제 독해 지문을 복습했을 때 잘 풀 수 있는
단어장·시험지·정답해설지를 만드는 것이다.

[전체 원칙]
1. 반드시 제공된 지문의 내용에 근거한다. 지문에 없는 사실을 만들어내지 않는다.
2. 지문 난이도와 어휘 밀도에 따라 뽑는 어휘 수를 자동 조절한다.
   쉬운 지문에서는 억지로 단어 수를 채우지 않고, 어려운 지문에서는 더 많이 뽑을 수 있다.
3. 단어뿐 아니라 독해에 중요한 phrasal verb, idiom, collocation, expression도 선택할 수 있다.
4. 고유명사나 지나치게 쉬운 단어는 원칙적으로 제외한다.
5. 목표 어휘는 지문 이해에 실제로 중요하거나 다른 독해에서도 재사용 가치가 높은 것을 우선한다.
6. 같은 어휘를 시험에서 불필요하게 반복하지 않는다.
7. 단어장 배열 순서와 시험 출제 순서는 같지 않게 한다.
8. 지문 5개 안팎이면 총 30문항 안팎을 기준으로 하되, 입력 분량과 가치에 따라 자동으로 줄이거나 늘린다.
9. 동의어/반의어는 해당 지문에서 사용된 뜻과 정확히 맞는 경우만 쓴다.
   정확한 관계가 없으면 반드시 "—"로 표시하고 억지로 만들지 않는다.
10. 모든 한국어 설명은 학생이 이해하기 쉽고 간결하게 쓴다.

[단어장]
각 항목 필수 정보:
- word/expression
- pronunciation: 한국 학생이 참고할 수 있는 IPA 또는 일반적인 영어 발음 표기
- pos: 품사
- meaning_ko: 지문에서의 정확한 한국어 뜻
- synonym: 정확한 문맥 동의어가 없으면 "—"
- antonym: 정확한 문맥 반의어가 없으면 "—"
- new_example: 원문과 다른 새 예문
- passage_label: 어느 지문/페이지에서 나온 것인지 짧은 표시
- importance: 1~5

새 예문은 시험 문제에 절대 그대로 재사용하지 않는다.

[시험 4유형]
시험지에는 유형 제목을 장황하게 쓰지 말고 숫자 1, 2, 3, 4만 표시한다.

1번 유형: 지문 복습형 주관식 + 첫 글자
- 원 지문의 핵심 사실/상황/인과/대조 중 하나 이상을 보존한다.
- 원문을 그대로 복사하고 단어 하나만 빈칸 처리하면 안 된다.
- 표현을 적당히 바꾸되 원문보다 문법을 지나치게 어렵게 만들지 않는다.
- 지문을 복습한 학생이면 상황을 알아볼 수 있어야 한다.
- 답의 첫 글자를 cue_first_letter에 제공한다.

2번 유형: 새 문맥 주관식 + 첫 글자
- 원 지문과 다른 완전히 새로운 상황을 사용한다.
- 단어장의 새 예문을 재사용하면 안 된다.
- 답의 첫 글자를 cue_first_letter에 제공한다.

3번 유형: 영영풀이 객관식
- 영어 정의를 보고 목표 어휘를 고르는 문제.
- 오답은 품사/난이도가 비슷하지만 뜻이 명확히 다른 어휘로 만든다.
- 4지선다를 기본으로 한다.

4번 유형: 동의어·반의어·어휘관계 객관식
- 실제로 정확한 관계가 있는 목표 어휘에 대해서만 출제한다.
- 관계가 애매하면 그 단어로 4번 유형을 억지로 만들지 않는다.
- 4지선다를 기본으로 한다.

[문항 비율]
기본 비율은 대략:
1번 40%
2번 30%
3번 15%
4번 15%
정수 문항 수에 맞게 가장 가깝게 배분한다.
단, 4번 유형의 정확한 재료가 부족하면 그 비중을 1~3번으로 자연스럽게 재분배한다.

[난이도]
- easier: 문맥 단서를 더 주고, 오답을 덜 헷갈리게 하며, 1번은 원문 상황에 더 가깝게.
- normal: 균형.
- harder: 문맥 단서를 조금 줄이고 오답을 더 그럴듯하게 하되,
  지문 문법 수준을 과도하게 넘지 않는다.

[정답·해설]
각 문항마다:
- answer
- meaning_ko
- explanation_ko
를 제공한다.
해설은 "왜 이 단어가 이 문맥에 맞는지" 또는 "왜 이 어휘관계가 맞는지"를 짧고 정확하게 설명한다.
문제 유형 이름을 정답지에서 반복 표기하지 않는다.

[출력 품질]
- source_note에는 입력 자료가 충분한지, 일부 페이지가 이미지라 텍스트 추출이 제한되었는지 등
  사용자가 알아야 할 점만 짧게 쓴다.
- 확실하지 않은 내용을 사실처럼 채우지 않는다.
"""


# ---------------------------------------------------------------------
# JSON Schema
# ---------------------------------------------------------------------

PACKAGE_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "title": {"type": "string"},
        "detected_level": {"type": "string"},
        "source_note": {"type": "string"},
        "vocabulary": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "id": {"type": "integer"},
                    "word": {"type": "string"},
                    "pronunciation": {"type": "string"},
                    "pos": {"type": "string"},
                    "meaning_ko": {"type": "string"},
                    "synonym": {"type": "string"},
                    "antonym": {"type": "string"},
                    "new_example": {"type": "string"},
                    "passage_label": {"type": "string"},
                    "importance": {"type": "integer", "minimum": 1, "maximum": 5},
                },
                "required": [
                    "id",
                    "word",
                    "pronunciation",
                    "pos",
                    "meaning_ko",
                    "synonym",
                    "antonym",
                    "new_example",
                    "passage_label",
                    "importance",
                ],
            },
        },
        "questions": {
            "type": "array",
            "items": {
                "type": "object",
                "additionalProperties": False,
                "properties": {
                    "number": {"type": "integer"},
                    "section": {"type": "integer", "enum": [1, 2, 3, 4]},
                    "target": {"type": "string"},
                    "cue_first_letter": {"type": "string"},
                    "prompt": {"type": "string"},
                    "choices": {
                        "type": "array",
                        "items": {"type": "string"},
                    },
                    "answer": {"type": "string"},
                    "meaning_ko": {"type": "string"},
                    "explanation_ko": {"type": "string"},
                },
                "required": [
                    "number",
                    "section",
                    "target",
                    "cue_first_letter",
                    "prompt",
                    "choices",
                    "answer",
                    "meaning_ko",
                    "explanation_ko",
                ],
            },
        },
    },
    "required": [
        "title",
        "detected_level",
        "source_note",
        "vocabulary",
        "questions",
    ],
}


QUESTION_SCHEMA: Dict[str, Any] = {
    "type": "object",
    "additionalProperties": False,
    "properties": {
        "number": {"type": "integer"},
        "section": {"type": "integer", "enum": [1, 2, 3, 4]},
        "target": {"type": "string"},
        "cue_first_letter": {"type": "string"},
        "prompt": {"type": "string"},
        "choices": {"type": "array", "items": {"type": "string"}},
        "answer": {"type": "string"},
        "meaning_ko": {"type": "string"},
        "explanation_ko": {"type": "string"},
    },
    "required": [
        "number",
        "section",
        "target",
        "cue_first_letter",
        "prompt",
        "choices",
        "answer",
        "meaning_ko",
        "explanation_ko",
    ],
}


# ---------------------------------------------------------------------
# PDF / 이미지 처리
# ---------------------------------------------------------------------

def parse_page_range(expr: str, total_pages: int) -> List[int]:
    """1-based 입력을 0-based 페이지 인덱스로 변환."""
    if not expr or not expr.strip():
        return list(range(total_pages))

    pages: set[int] = set()

    for token in expr.split(","):
        token = token.strip()
        if not token:
            continue

        if "-" in token:
            a_s, b_s = token.split("-", 1)
            a, b = int(a_s.strip()), int(b_s.strip())
            start, end = min(a, b), max(a, b)
            for p in range(start, end + 1):
                pages.add(p - 1)
        else:
            pages.add(int(token) - 1)

    valid = sorted(p for p in pages if 0 <= p < total_pages)

    if not valid:
        raise ValueError(
            f"입력한 페이지 범위가 PDF의 실제 범위(1~{total_pages})와 맞지 않습니다."
        )

    return valid


def image_to_data_url(png_bytes: bytes, mime: str = "image/png") -> str:
    encoded = base64.b64encode(png_bytes).decode("ascii")
    return f"data:{mime};base64,{encoded}"


def pdf_to_source_parts(
    pdf_bytes: bytes,
    page_range: str,
    file_label: str,
) -> Tuple[List[Dict[str, Any]], str]:
    """
    텍스트가 충분한 페이지는 텍스트로,
    텍스트가 거의 없는 페이지는 렌더링 이미지로 모델에 전달.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    selected = parse_page_range(page_range, len(doc))

    parts: List[Dict[str, Any]] = []
    text_pages = 0
    image_pages = 0

    for idx in selected:
        page = doc[idx]
        text = page.get_text("text").strip()

        # 텍스트가 어느 정도 있으면 비용/속도상 텍스트 우선
        if len(text) >= 120:
            text_pages += 1
            parts.append(
                {
                    "type": "input_text",
                    "text": f"\n--- {file_label} / PDF page {idx + 1} ---\n{text}",
                }
            )
        else:
            image_pages += 1
            pix = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            png = pix.tobytes("png")
            parts.append(
                {
                    "type": "input_text",
                    "text": f"\n--- {file_label} / PDF page {idx + 1} (image page) ---",
                }
            )
            parts.append(
                {
                    "type": "input_image",
                    "image_url": image_to_data_url(png),
                    "detail": "high",
                }
            )

    doc.close()

    note = (
        f"{file_label}: 선택 {len(selected)}쪽 "
        f"(텍스트 {text_pages}쪽, 이미지 분석 {image_pages}쪽)"
    )
    return parts, note


def uploaded_files_to_source_parts(
    uploaded_files,
    page_range: str,
) -> Tuple[List[Dict[str, Any]], str]:
    all_parts: List[Dict[str, Any]] = []
    notes: List[str] = []

    for f in uploaded_files:
        raw = f.getvalue()
        name = f.name or "uploaded"
        lower = name.lower()
        mime = f.type or ""

        if lower.endswith(".pdf") or mime == "application/pdf":
            parts, note = pdf_to_source_parts(raw, page_range, name)
            all_parts.extend(parts)
            notes.append(note)
        elif lower.endswith((".png", ".jpg", ".jpeg", ".webp")) or mime.startswith("image/"):
            actual_mime = mime if mime.startswith("image/") else "image/jpeg"
            all_parts.append(
                {
                    "type": "input_text",
                    "text": f"\n--- image file: {name} ---",
                }
            )
            all_parts.append(
                {
                    "type": "input_image",
                    "image_url": image_to_data_url(raw, actual_mime),
                    "detail": "high",
                }
            )
            notes.append(f"{name}: 이미지 1개")
        else:
            raise ValueError(
                f"{name}: 지원하지 않는 파일 형식입니다. PDF/JPG/JPEG/PNG/WEBP만 올려 주세요."
            )

    return all_parts, " | ".join(notes)


# ---------------------------------------------------------------------
# OpenAI 응답 처리
# ---------------------------------------------------------------------

def extract_json_text(raw: str) -> Dict[str, Any]:
    text = raw.strip()

    # ```json ... ``` 대응
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

    return json.loads(text)


def responses_json(
    *,
    content: List[Dict[str, Any]],
    schema: Dict[str, Any],
    schema_name: str,
    model: str,
) -> Dict[str, Any]:
    """
    Structured Output을 우선 사용하고,
    SDK/모델 조합에서 포맷 인자가 지원되지 않는 경우 JSON-only 지시로 한 번 재시도.
    """
    try:
        resp = client.responses.create(
            model=model,
            instructions=SYSTEM_MANUAL,
            input=[{"role": "user", "content": content}],
            text={
                "format": {
                    "type": "json_schema",
                    "name": schema_name,
                    "strict": True,
                    "schema": schema,
                }
            },
        )
        return extract_json_text(resp.output_text)
    except Exception as first_error:
        # 포맷 호환성 문제에 대비한 fallback.
        fallback_instruction = {
            "type": "input_text",
            "text": (
                "\n반드시 JSON만 출력하세요. Markdown code fence를 쓰지 마세요.\n"
                "다음 JSON Schema의 필드 구조를 정확히 따르세요:\n"
                + json.dumps(schema, ensure_ascii=False)
            ),
        }

        try:
            resp = client.responses.create(
                model=model,
                instructions=SYSTEM_MANUAL,
                input=[{"role": "user", "content": content + [fallback_instruction]}],
            )
            return extract_json_text(resp.output_text)
        except Exception:
            raise first_error


def desired_question_hint(vocab_count: int, source_count_hint: int) -> str:
    # AI가 최종 판단하되 과도한 filler를 막기 위한 부드러운 힌트.
    if vocab_count <= 0:
        return "문항 수는 입력 분량에 맞게 자동 결정"
    approx = min(40, max(12, round(vocab_count * 0.9)))
    return f"총 문항 수는 약 {approx}개를 참고하되, 가치 없는 filler는 만들지 말 것"


def generate_package(
    title: str,
    difficulty: str,
    uploaded_files,
    page_range: str,
    model: str,
) -> Dict[str, Any]:
    source_parts, local_note = uploaded_files_to_source_parts(
        uploaded_files,
        page_range,
    )

    task = {
        "type": "input_text",
        "text": f"""
[작업]
시험지 제목: {title or "Reading Vocabulary Review"}
난이도: {difficulty}

업로드한 독해 자료를 분석하여 아래를 한 세트로 생성하세요.
1) 단어장
2) 1·2·3·4 네 유형의 시험지
3) 정답·해설지

중요:
- 지문별 어휘 수와 전체 문항 수는 자동 조절.
- 약 5개 지문 분량이면 30문항 안팎을 참고.
- 단어장 새 예문과 2번 새 문맥 문제는 서로 다른 문장을 사용.
- 1번은 원문 상황을 알아볼 수 있으나 원문 그대로의 빈칸 문제가 아니어야 함.
- 4번은 정확한 동의어/반의어/관계가 있는 경우에만 사용.
- 시험 어휘 순서는 단어장 순서와 다르게 섞을 것.
- choices는 1·2번에서는 빈 배열, 3·4번에서는 4개를 권장.
- cue_first_letter는 1·2번에서 답의 첫 글자, 3·4번에서는 빈 문자열.

로컬 전처리 참고: {local_note}
""",
    }

    data = responses_json(
        content=source_parts + [task],
        schema=PACKAGE_SCHEMA,
        schema_name="reading_vocabulary_package",
        model=model,
    )

    # 화면 제목을 사용자가 입력한 값으로 우선 고정
    if title.strip():
        data["title"] = title.strip()

    return validate_and_shuffle(data)


def regenerate_question(
    data: Dict[str, Any],
    question_number: int,
    difficulty: str,
    model: str,
) -> Dict[str, Any]:
    original = next(
        (q for q in data["questions"] if q["number"] == question_number),
        None,
    )
    if not original:
        raise ValueError("선택한 문항을 찾을 수 없습니다.")

    vocab_item = next(
        (v for v in data["vocabulary"] if v["word"].lower() == original["target"].lower()),
        None,
    )

    context = {
        "type": "input_text",
        "text": f"""
다음 기존 문항 하나를 같은 target과 같은 section으로 다시 작성하세요.
난이도: {difficulty}

기존 문항:
{json.dumps(original, ensure_ascii=False)}

관련 단어장 정보:
{json.dumps(vocab_item or {}, ensure_ascii=False)}

규칙:
- section과 target은 절대 바꾸지 말 것.
- 1번이면 기존 문항에 담긴 원지문 사실/상황을 보존하면서 표현과 단서만 조절.
- 2번이면 기존 문항 및 단어장 예문과 다른 새 상황.
- 3·4번 객관식이면 choices를 4개로.
- 정답과 한국어 뜻, 해설도 다시 정리.
- cue_first_letter는 1·2번만 사용.
""",
    }

    new_q = responses_json(
        content=[context],
        schema=QUESTION_SCHEMA,
        schema_name="reading_vocabulary_question",
        model=model,
    )

    # 핵심 식별자는 원본을 보존
    new_q["number"] = original["number"]
    new_q["section"] = original["section"]
    new_q["target"] = original["target"]

    out = deepcopy(data)
    out["questions"] = [
        new_q if q["number"] == question_number else q
        for q in out["questions"]
    ]
    return out


# ---------------------------------------------------------------------
# 데이터 검증 / 섞기
# ---------------------------------------------------------------------

def normalize_choice(choice: str, idx: int) -> str:
    # 모델이 이미 ① 등을 붙여도 중복시키지 않음
    c = str(choice).strip()
    if re.match(r"^[①②③④]\s*", c):
        return c
    labels = ["①", "②", "③", "④"]
    prefix = labels[idx] if idx < len(labels) else f"{idx + 1}."
    return f"{prefix} {c}"


def validate_and_shuffle(data: Dict[str, Any]) -> Dict[str, Any]:
    out = deepcopy(data)

    vocab = out.get("vocabulary", [])
    questions = out.get("questions", [])

    # synonym / antonym 비어 있으면 —
    for v in vocab:
        v["synonym"] = (v.get("synonym") or "—").strip()
        v["antonym"] = (v.get("antonym") or "—").strip()

    # section별 순서 랜덤
    rng = random.SystemRandom()
    groups = {1: [], 2: [], 3: [], 4: []}

    for q in questions:
        sec = int(q.get("section", 1))
        q["section"] = sec

        if sec in (1, 2):
            q["choices"] = []
            if not q.get("cue_first_letter"):
                ans = q.get("answer", "").strip()
                q["cue_first_letter"] = ans[:1] if ans else ""
        else:
            q["cue_first_letter"] = ""
            q["choices"] = [
                normalize_choice(c, i)
                for i, c in enumerate(q.get("choices", [])[:4])
            ]

        groups.setdefault(sec, []).append(q)

    final_questions: List[Dict[str, Any]] = []
    for sec in (1, 2, 3, 4):
        rng.shuffle(groups.get(sec, []))
        final_questions.extend(groups.get(sec, []))

    for i, q in enumerate(final_questions, 1):
        q["number"] = i

    out["questions"] = final_questions
    return out


def reshuffle_only(data: Dict[str, Any]) -> Dict[str, Any]:
    out = deepcopy(data)
    rng = random.SystemRandom()

    groups = {1: [], 2: [], 3: [], 4: []}
    for q in out["questions"]:
        groups[q["section"]].append(q)

    qs = []
    for sec in (1, 2, 3, 4):
        rng.shuffle(groups[sec])
        qs.extend(groups[sec])

    for i, q in enumerate(qs, 1):
        q["number"] = i

    out["questions"] = qs
    return out


# ---------------------------------------------------------------------
# Word 문서 공통 서식
# ---------------------------------------------------------------------

def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for m, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_page_border(section):
    """
    본문 테두리는 page보다 text 기준으로 안쪽에 두고,
    footer는 더 아래에 배치하여 motto가 시각적으로 테두리 밖에 오도록 함.
    """
    sect_pr = section._sectPr

    old = sect_pr.find(qn("w:pgBorders"))
    if old is not None:
        sect_pr.remove(old)

    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "text")

    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8")
        tag.set(qn("w:space"), "8")
        tag.set(qn("w:color"), "666666")
        pg_borders.append(tag)

    sect_pr.append(pg_borders)


def add_footer_motto(section):
    section.footer_distance = Cm(0.25)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Hard work pays off.")
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


def base_document() -> Document:
    doc = Document()
    sec = doc.sections[0]

    sec.top_margin = Cm(1.35)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.15)
    sec.right_margin = Cm(1.15)

    set_page_border(sec)
    add_footer_motto(sec)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)

    return doc


def add_document_heading(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)

    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(6)

    r2 = p2.add_run(subtitle)
    r2.bold = True
    r2.font.size = Pt(12)


def docx_to_bytes(doc: Document) -> bytes:
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


# ---------------------------------------------------------------------
# 단어장 Word
# ---------------------------------------------------------------------

def build_vocabulary_docx(data: Dict[str, Any]) -> bytes:
    doc = base_document()
    add_document_heading(doc, data["title"], "Vocabulary List")

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    widths = [Cm(0.9), Cm(3.0), Cm(2.5), Cm(3.0), Cm(3.1), Cm(6.7)]
    headers = ["No.", "Word / Expression", "발음 · 품사", "뜻", "Syn. / Ant.", "New Example"]

    hdr = table.rows[0]
    set_repeat_table_header(hdr)

    for i, (cell, header) in enumerate(zip(hdr.cells, headers)):
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "EDEDED")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(10)

    for index, v in enumerate(data["vocabulary"], 1):
        cells = table.add_row().cells

        relation = []
        if v["synonym"] != "—":
            relation.append("S: " + v["synonym"])
        if v["antonym"] != "—":
            relation.append("A: " + v["antonym"])

        values = [
            str(index),
            v["word"],
            f'{v["pronunciation"]}\n{v["pos"]}',
            v["meaning_ko"],
            "\n".join(relation) if relation else "—",
            v["new_example"],
        ]

        for i, (cell, text) in enumerate(zip(cells, values)):
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, bottom=80)

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05

            if i == 1:
                # 큰 영어 글씨
                r = p.add_run(text)
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(12.5)
            elif i == 5:
                # 예문 칸을 가장 넓고 읽기 쉽게
                r = p.add_run(text)
                r.font.name = "Arial"
                r.font.size = Pt(10.7)
            else:
                r = p.add_run(text)
                r.font.size = Pt(9.8)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    r = note.add_run(
        "※ 새 예문은 학습용 문장이며 시험 문제에 그대로 재사용하지 않습니다. "
        "정확한 동의어·반의어가 없으면 —로 표시합니다."
    )
    r.font.size = Pt(8.8)

    return docx_to_bytes(doc)


# ---------------------------------------------------------------------
# 시험지 Word
# ---------------------------------------------------------------------

def format_question_for_test(q: Dict[str, Any]) -> str:
    prompt = q["prompt"].strip()

    if q["section"] in (1, 2):
        cue = q.get("cue_first_letter", "").strip()
        if cue:
            return f"{prompt}\n({cue}__________)"
        return prompt

    choices = q.get("choices", [])
    if choices:
        return prompt + "\n" + "   ".join(choices)

    return prompt


def add_question_block(cell, q: Dict[str, Any]):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.35

    number_run = p.add_run(f'{q["number"]}. ')
    number_run.bold = True
    number_run.font.size = Pt(12.4)

    text_run = p.add_run(format_question_for_test(q))
    text_run.font.name = "Arial"
    text_run.font.size = Pt(12.4)


def build_test_docx(data: Dict[str, Any]) -> bytes:
    doc = base_document()
    add_document_heading(doc, data["title"], "Vocabulary Review Test")

    info = doc.add_table(rows=1, cols=3)
    info.style = "Table Grid"
    info.autofit = False
    info.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = [
        "학원/반: ____________________",
        "이름: ____________________",
        "점수: ______ / ______",
    ]

    for i, text in enumerate(labels):
        cell = info.cell(0, i)
        cell.width = Cm(6.0)
        set_cell_margins(cell, top=100, bottom=100)
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.size = Pt(11)

    directions = doc.add_paragraph()
    directions.paragraph_format.space_before = Pt(4)
    directions.paragraph_format.space_after = Pt(4)
    r = directions.add_run(
        "1·2: 첫 글자를 참고하여 알맞은 단어를 쓰시오.   "
        "3·4: 가장 알맞은 답을 고르시오."
    )
    r.bold = True
    r.font.size = Pt(10.5)

    # 섹션은 숫자만 표시.
    for section_no in (1, 2, 3, 4):
        section_questions = [
            q for q in data["questions"] if q["section"] == section_no
        ]

        if not section_questions:
            continue

        section_title = doc.add_paragraph()
        section_title.paragraph_format.space_before = Pt(5)
        section_title.paragraph_format.space_after = Pt(2)
        rr = section_title.add_run(str(section_no))
        rr.bold = True
        rr.font.size = Pt(15)

        # 문항을 두 열로 분배
        two_col = doc.add_table(rows=1, cols=2)
        two_col.autofit = False
        two_col.alignment = WD_TABLE_ALIGNMENT.CENTER

        left, right = two_col.rows[0].cells
        left.width = Cm(9.0)
        right.width = Cm(9.0)

        for c in (left, right):
            set_cell_margins(c, top=20, bottom=20, start=120, end=120)
            # 2단 칸의 바깥 테두리는 눈에 띄지 않게 제거
            tc_pr = c._tc.get_or_add_tcPr()
            tc_borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "nil")
                tc_borders.append(e)
            tc_pr.append(tc_borders)

        midpoint = (len(section_questions) + 1) // 2
        left_qs = section_questions[:midpoint]
        right_qs = section_questions[midpoint:]

        for q in left_qs:
            add_question_block(left, q)

        for q in right_qs:
            add_question_block(right, q)

    return docx_to_bytes(doc)


# ---------------------------------------------------------------------
# 정답·해설지 Word
# ---------------------------------------------------------------------

def build_answer_docx(data: Dict[str, Any]) -> bytes:
    doc = base_document()
    add_document_heading(doc, data["title"], "Answer & Explanation")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    # 번호·정답 좁게, 뜻 조금 넓게, 해설 최대
    widths = [Cm(1.1), Cm(2.8), Cm(4.0), Cm(10.0)]
    headers = ["No.", "정답", "뜻", "문제 해설"]

    hdr = table.rows[0]
    set_repeat_table_header(hdr)

    for i, (cell, header) in enumerate(zip(hdr.cells, headers)):
        cell.width = widths[i]
        set_cell_shading(cell, "EDEDED")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(10)

    for q in data["questions"]:
        values = [
            str(q["number"]),
            q["answer"],
            q["meaning_ko"],
            q["explanation_ko"],
        ]

        cells = table.add_row().cells

        for i, (cell, text) in enumerate(zip(cells, values)):
            cell.width = widths[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=85, bottom=85)

            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(text)
            r.font.size = Pt(10.3 if i < 3 else 10.5)

            if i in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return docx_to_bytes(doc)


# ---------------------------------------------------------------------
# 화면
# ---------------------------------------------------------------------

st.title("📘 Reading Vocabulary Maker")
st.caption("독해 복습형 단어장 · 시험지 · 정답해설지 생성")

with st.container(border=True):
    col1, col2, col3 = st.columns([2.2, 1.25, 1.0])

    with col1:
        title = st.text_input(
            "시험지 제목",
            placeholder="예: Hackers TEPS Reading - Chapter 05",
        )

    with col2:
        page_range = st.text_input(
            "PDF 페이지 범위",
            placeholder="예: 94-113 또는 94,96-100",
            help="여러 PDF를 올리면 같은 페이지 범위를 각 PDF에 적용합니다.",
        )

    with col3:
        difficulty_label = st.selectbox(
            "난이도",
            ["더 쉽게", "기본", "더 어렵게"],
            index=1,
        )

    # Tablet/Android compatibility: use a single-file uploader without a browser-side
    # extension filter. Some Android/Samsung file pickers return generic MIME types
    # or behave unreliably with multi-select + accept filters. We validate the actual
    # filename/MIME later in uploaded_files_to_source_parts().
    uploaded_file = st.file_uploader(
        "독해 자료 PDF / 이미지",
        accept_multiple_files=False,
        key="source_file_single",
        help="PDF/JPG/JPEG/PNG/WEBP 1개를 선택하세요. 태블릿 호환성을 위해 한 번에 1개씩 받습니다.",
    )

    if uploaded_file is not None:
        st.success(f"선택됨: {uploaded_file.name} ({uploaded_file.size / 1024 / 1024:.1f} MB)")
        uploaded_files = [uploaded_file]
    else:
        uploaded_files = []

with st.expander("생성 규칙 확인", expanded=False):
    st.write(
        "• 지문 난이도·어휘 밀도에 따라 단어 수 자동 조절\n"
        "• 단어장: 단어/발음/품사/뜻/정확한 동의어·반의어/새 예문\n"
        "• 1: 지문 복습형 주관식 + 첫 글자\n"
        "• 2: 새 문맥 주관식 + 첫 글자\n"
        "• 3: 영영풀이 객관식\n"
        "• 4: 동의어·반의어·어휘관계 객관식\n"
        "• 단어장 예문은 시험에 재사용하지 않음\n"
        "• 시험 순서는 단어장 순서와 다르게 섞음"
    )

difficulty_map = {
    "더 쉽게": "easier",
    "기본": "normal",
    "더 어렵게": "harder",
}

if st.button(
    "분석하고 단어장 · 시험지 · 해설지 만들기",
    type="primary",
    use_container_width=True,
):
    if not uploaded_files:
        st.warning("PDF 또는 이미지 파일을 하나 이상 올려 주세요.")
    else:
        with st.spinner(
            "지문을 분석하고 있습니다. 분량이 많으면 시간이 조금 걸릴 수 있습니다."
        ):
            try:
                result = generate_package(
                    title=title,
                    difficulty=difficulty_map[difficulty_label],
                    uploaded_files=uploaded_files,
                    page_range=page_range,
                    model=DEFAULT_MODEL,
                )
                st.session_state["vocab_result"] = result
                st.success("생성이 완료되었습니다.")
            except Exception as exc:
                st.error("생성 중 오류가 발생했습니다.")
                st.exception(exc)

data = st.session_state.get("vocab_result")

if data:
    st.info(
        f"추정 지문 수준: {data.get('detected_level', '—')}"
        + (
            f"\n\n{data.get('source_note')}"
            if data.get("source_note")
            else ""
        )
    )

    top1, top2, top3 = st.columns([1.2, 1.2, 3.6])

    with top1:
        if st.button("시험 순서 다시 섞기", use_container_width=True):
            st.session_state["vocab_result"] = reshuffle_only(data)
            st.rerun()

    tab_vocab, tab_test, tab_answer = st.tabs(
        ["단어장", "시험지", "정답·해설"]
    )

    with tab_vocab:
        st.write(f"선정 어휘: **{len(data['vocabulary'])}개**")
        vocab_rows = []

        for i, v in enumerate(data["vocabulary"], 1):
            vocab_rows.append(
                {
                    "No.": i,
                    "Word / Expression": v["word"],
                    "발음": v["pronunciation"],
                    "품사": v["pos"],
                    "뜻": v["meaning_ko"],
                    "Synonym": v["synonym"],
                    "Antonym": v["antonym"],
                    "New Example": v["new_example"],
                    "출처": v["passage_label"],
                }
            )

        st.dataframe(
            vocab_rows,
            use_container_width=True,
            hide_index=True,
        )

    with tab_test:
        st.write(f"총 문항: **{len(data['questions'])}문항**")

        for sec in (1, 2, 3, 4):
            qs = [q for q in data["questions"] if q["section"] == sec]
            if not qs:
                continue

            st.subheader(str(sec))

            for q in qs:
                text = q["prompt"]

                if sec in (1, 2) and q.get("cue_first_letter"):
                    text += f"\n\n({q['cue_first_letter']}__________)"

                st.markdown(f"**{q['number']}.** {text}")

                if q.get("choices"):
                    st.write("　".join(q["choices"]))

                b1, b2, _ = st.columns([1.15, 1.15, 4.7])

                with b1:
                    if st.button(
                        "이 문제 쉽게",
                        key=f"easy_{q['number']}",
                        use_container_width=True,
                    ):
                        with st.spinner("문항을 쉽게 다시 만드는 중..."):
                            st.session_state["vocab_result"] = regenerate_question(
                                data,
                                q["number"],
                                "easier",
                                DEFAULT_MODEL,
                            )
                        st.rerun()

                with b2:
                    if st.button(
                        "이 문제 어렵게",
                        key=f"hard_{q['number']}",
                        use_container_width=True,
                    ):
                        with st.spinner("문항을 어렵게 다시 만드는 중..."):
                            st.session_state["vocab_result"] = regenerate_question(
                                data,
                                q["number"],
                                "harder",
                                DEFAULT_MODEL,
                            )
                        st.rerun()

                st.divider()

    with tab_answer:
        rows = []

        for q in data["questions"]:
            rows.append(
                {
                    "No.": q["number"],
                    "정답": q["answer"],
                    "뜻": q["meaning_ko"],
                    "문제 해설": q["explanation_ko"],
                }
            )

        st.dataframe(
            rows,
            use_container_width=True,
            hide_index=True,
        )

    st.subheader("Word 파일 다운로드")

    vocab_doc = build_vocabulary_docx(data)
    test_doc = build_test_docx(data)
    answer_doc = build_answer_docx(data)

    dl1, dl2, dl3 = st.columns(3)

    mime_docx = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    with dl1:
        st.download_button(
            "단어장 Word",
            data=vocab_doc,
            file_name="Vocabulary_List.docx",
            mime=mime_docx,
            use_container_width=True,
        )

    with dl2:
        st.download_button(
            "시험지 Word",
            data=test_doc,
            file_name="Vocabulary_Test.docx",
            mime=mime_docx,
            use_container_width=True,
        )

    with dl3:
        st.download_button(
            "정답·해설 Word",
            data=answer_doc,
            file_name="Answer_Explanation.docx",
            mime=mime_docx,
            use_container_width=True,
        )

    st.caption("Hard work pays off.")

from __future__ import annotations

import base64
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Literal

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel, Field, field_validator
from pypdf import PdfReader, PdfWriter

APP_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = APP_DIR / "templates"
LIST_TEMPLATE = TEMPLATE_DIR / "Vocabulary_List_Template.docx"
TEST_TEMPLATE = TEMPLATE_DIR / "Vocabulary_Test_Template.docx"
ANSWER_TEMPLATE = TEMPLATE_DIR / "Answer_Explanation_Template.docx"

MODEL_DEFAULT = os.getenv("OPENAI_MODEL", "gpt-5.6")

# -----------------------------
# Structured data models
# -----------------------------

class VocabularyItem(BaseModel):
    word: str
    pronunciation: str
    pos: str
    meaning_ko: str
    syn_ant: str = "—"
    new_example: str
    source_evidence: str = ""

    @field_validator("word", "pronunciation", "pos", "meaning_ko", "new_example")
    @classmethod
    def non_empty(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("empty field")
        return v


class VocabularyAnalysis(BaseModel):
    chapter_title: str
    unit_title: str
    source_summary_ko: str
    vocabulary: list[VocabularyItem]


class QuestionDraft(BaseModel):
    number: int = Field(ge=1, le=30)
    section: Literal[1, 2, 3, 4]
    stem: str
    target_word: str
    meaning_ko: str
    explanation_ko: str
    choices: list[str] = Field(default_factory=list)
    correct_choice: int | None = None
    source_evidence: str = ""


class TestBundle(BaseModel):
    questions: list[QuestionDraft]


# -----------------------------
# Prompt contracts
# -----------------------------

VOCAB_SYSTEM = """You are an expert Korean TEPS reading instructor and vocabulary editor.
Use ONLY the supplied source pages as the basis for vocabulary selection.
Do not add facts that are not supported by the source.
Select only genuinely useful learning vocabulary and expressions. Do NOT target a fixed count such as 50.
Avoid padding with easy or low-value words merely to increase the count.
Include important phrases and phrasal verbs when educationally useful.
For each item:
- word: exact useful headword/expression
- pronunciation: IPA only, enclosed in slashes
- pos: concise label such as n., v., adj., adv., phr., phr.v., n./v.
- meaning_ko: concise Korean meaning appropriate to the source context
- syn_ant: use 'S: ...' and/or 'A: ...'; if no precise useful match, use '—'
- new_example: one NEW learning sentence, not copied from the source
- source_evidence: short source phrase or description showing where/why the item was selected
Do not reuse a source sentence verbatim as the new example.
"""

TEST_SYSTEM = """You are an expert Korean TEPS vocabulary test writer.
Use ONLY the supplied source pages and the APPROVED vocabulary list.
Create exactly 30 questions with this fixed structure:
1. Reading Review: 1-12. Source-grounded review of actual passage content/context. Fill-in-the-blank. stem must contain exactly one token [[BLANK]].
2. Vocabulary Transfer: 13-21. New context, not copied from the Vocabulary List New Example and not copied from source. Fill-in-the-blank. stem must contain exactly one token [[BLANK]].
3. English Definition: 22-26. Four-choice English definitions. choices must contain exactly four plain-text options; correct_choice 1-4.
4. Vocabulary Relations: 27-30. Four-choice synonym/antonym/expression relationship. choices exactly four; correct_choice 1-4.
Rules:
- target_word must be exactly one entry from the approved vocabulary list (case-insensitive exact headword/expression).
- For sections 1 and 2, do NOT write first-letter underscores yourself; use [[BLANK]] and the app will generate the hint.
- For sections 3 and 4, do not include ①②③④ inside choices; provide clean choice strings only.
- Every question needs a concise Korean meaning and a real explanation of WHY the answer is correct.
- Reading Review 1-12 must be tied to actual source content; source_evidence must briefly identify the source fact/context.
- Do not reuse any Vocabulary List new_example sentence.
- Avoid ambiguous items; each multiple-choice question must have exactly one defensible answer.
"""

# -----------------------------
# Source handling
# -----------------------------

def parse_page_range(text: str, total_pages: int) -> list[int]:
    text = (text or "").strip()
    if not text:
        return list(range(1, total_pages + 1))
    pages: set[int] = set()
    for token in re.split(r"\s*,\s*", text):
        if not token:
            continue
        if "-" in token:
            a, b = token.split("-", 1)
            a, b = int(a), int(b)
            if a > b:
                a, b = b, a
            pages.update(range(a, b + 1))
        else:
            pages.add(int(token))
    bad = [p for p in pages if p < 1 or p > total_pages]
    if bad:
        raise ValueError(f"페이지 범위가 PDF 범위를 벗어났습니다: {bad[:5]} (전체 {total_pages}쪽)")
    return sorted(pages)


def subset_pdf(pdf_bytes: bytes, page_range: str) -> tuple[bytes, list[int], int]:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    total = len(reader.pages)
    selected = parse_page_range(page_range, total)
    writer = PdfWriter()
    for p in selected:
        writer.add_page(reader.pages[p - 1])
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue(), selected, total


def data_uri(mime: str, data: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(data).decode('ascii')}"


@dataclass
class PreparedSource:
    content_items: list[dict]
    note: str


def prepare_sources(uploaded_files, page_range: str) -> PreparedSource:
    content: list[dict] = []
    notes: list[str] = []
    pdf_count = 0
    total_payload_bytes = 0
    for f in uploaded_files:
        raw = f.getvalue()
        mime = f.type or "application/octet-stream"
        name = f.name
        if name.lower().endswith(".pdf") or mime == "application/pdf":
            pdf_count += 1
            if page_range and len(uploaded_files) > 1:
                raise ValueError("페이지 범위 지정은 PDF 1개 업로드일 때만 지원합니다.")
            if page_range:
                raw, selected, total = subset_pdf(raw, page_range)
                notes.append(f"{name}: PDF {total}쪽 중 {selected[0]}-{selected[-1]}쪽 사용")
            else:
                notes.append(f"{name}: 전체 PDF 사용")
            total_payload_bytes += len(raw)
            content.append({
                "type": "input_file",
                "filename": name,
                "file_data": data_uri("application/pdf", raw),
                "detail": "high",
            })
        elif mime.startswith("image/") or name.lower().endswith((".jpg", ".jpeg", ".png", ".webp")):
            notes.append(f"{name}: 이미지 사용")
            total_payload_bytes += len(raw)
            content.append({
                "type": "input_image",
                "image_url": data_uri(mime if mime.startswith("image/") else "image/jpeg", raw),
                "detail": "high",
            })
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {name}")
    if not content:
        raise ValueError("업로드된 자료가 없습니다.")
    if total_payload_bytes >= 50 * 1024 * 1024:
        mb = total_payload_bytes / (1024 * 1024)
        raise ValueError(f"OpenAI 파일 입력 한도 때문에 분석에 보낼 자료는 합계 50MB 미만이어야 합니다. 현재 {mb:.1f}MB입니다. 큰 PDF는 필요한 페이지 범위를 지정해 주세요.")
    return PreparedSource(content, " | ".join(notes))


# -----------------------------
# OpenAI calls
# -----------------------------

def get_client(api_key: str | None = None):
    from openai import OpenAI
    key = api_key or os.getenv("OPENAI_API_KEY", "")
    if not key:
        raise RuntimeError("OPENAI_API_KEY가 없습니다.")
    return OpenAI(api_key=key)


def analyze_vocabulary(prepared: PreparedSource, title: str, unit_title: str, model: str, api_key: str | None = None) -> VocabularyAnalysis:
    client = get_client(api_key)
    user_text = f"""자료를 분석해 Vocabulary List 후보를 만들어 주세요.
문서 제목: {title or 'Reading Vocabulary Maker'}
단원명: {unit_title or 'Vocabulary List'}
중요: 단어 수는 고정하지 말고, 실제 학습 가치가 있는 항목만 선택하세요.
"""
    response = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": VOCAB_SYSTEM},
            {"role": "user", "content": [{"type": "input_text", "text": user_text}] + prepared.content_items},
        ],
        text_format=VocabularyAnalysis,
    )
    if response.output_parsed is None:
        raise RuntimeError("어휘 분석 결과를 구조화하지 못했습니다.")
    return response.output_parsed


def generate_test(prepared: PreparedSource, vocab: list[VocabularyItem], title: str, unit_title: str, model: str, api_key: str | None = None) -> TestBundle:
    client = get_client(api_key)
    approved = [v.model_dump(exclude={"source_evidence"}) for v in vocab]
    user_text = f"""다음 승인된 Vocabulary List만 정답 후보로 사용해 30문항 시험과 해설을 만들어 주세요.
문서 제목: {title}
단원명: {unit_title}
APPROVED VOCABULARY JSON:
{json.dumps(approved, ensure_ascii=False)}
"""
    response = client.responses.parse(
        model=model,
        input=[
            {"role": "system", "content": TEST_SYSTEM},
            {"role": "user", "content": [{"type": "input_text", "text": user_text}] + prepared.content_items},
        ],
        text_format=TestBundle,
    )
    if response.output_parsed is None:
        raise RuntimeError("시험 생성 결과를 구조화하지 못했습니다.")
    return response.output_parsed


# -----------------------------
# Validation / QA
# -----------------------------

def norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", s.lower()).strip()


def validate_vocab(vocab: list[VocabularyItem]) -> list[str]:
    issues: list[str] = []
    if not vocab:
        issues.append("Vocabulary List가 비어 있습니다.")
        return issues
    words = [norm(v.word) for v in vocab]
    dup = sorted({w for w in words if words.count(w) > 1 and w})
    if dup:
        issues.append(f"중복 어휘가 있습니다: {', '.join(dup[:8])}")
    for i, v in enumerate(vocab, 1):
        if not (v.pronunciation.startswith("/") and v.pronunciation.endswith("/")):
            issues.append(f"{i}번 {v.word}: IPA가 /.../ 형식이 아닙니다.")
        if not v.pos.strip():
            issues.append(f"{i}번 {v.word}: 품사가 없습니다.")
        if len(v.new_example.split()) < 4:
            issues.append(f"{i}번 {v.word}: New Example이 너무 짧습니다.")
    return issues


def build_hint(answer: str) -> str:
    # 각 단어의 첫 글자를 살리고 나머지를 밑줄로 표시. 하이픈은 유지.
    chunks = re.split(r"(\s+|-)", answer.strip())
    out = []
    for chunk in chunks:
        if not chunk or chunk.isspace() or chunk == "-":
            out.append(chunk)
            continue
        letters = re.sub(r"[^A-Za-z]", "", chunk)
        if not letters:
            out.append(chunk)
            continue
        first = letters[0]
        out.append(first + "_" * max(3, len(letters) - 1))
    return "".join(out)


def render_question(q: QuestionDraft) -> tuple[str, str]:
    if q.section in (1, 2):
        prompt = q.stem.replace("[[BLANK]]", build_hint(q.target_word))
        return prompt, q.target_word
    circled = ["①", "②", "③", "④"]
    choices = "   ".join(f"{circled[i]} {c}" for i, c in enumerate(q.choices))
    prompt = f"{q.stem}\n    {choices}"
    answer = f"{circled[(q.correct_choice or 1)-1]} {q.choices[(q.correct_choice or 1)-1]}"
    return prompt, answer


def validate_test(bundle: TestBundle, vocab: list[VocabularyItem]) -> list[str]:
    issues: list[str] = []
    qs = bundle.questions
    if len(qs) != 30:
        issues.append(f"시험 문항 수가 30개가 아닙니다: {len(qs)}개")
        return issues
    nums = [q.number for q in qs]
    if sorted(nums) != list(range(1, 31)):
        issues.append("문항 번호가 1~30 연속이 아닙니다.")
    expected_section = {n: 1 if n <= 12 else 2 if n <= 21 else 3 if n <= 26 else 4 for n in range(1, 31)}
    for q in qs:
        if q.section != expected_section[q.number]:
            issues.append(f"{q.number}번 섹션 배치 오류")
    approved = {norm(v.word): v for v in vocab}
    examples = [norm(v.new_example) for v in vocab]
    stems_seen: set[str] = set()
    for q in qs:
        tw = norm(q.target_word)
        if tw not in approved:
            issues.append(f"{q.number}번 target_word가 승인 Vocabulary List에 없습니다: {q.target_word}")
        sn = norm(q.stem)
        if sn in stems_seen:
            issues.append(f"{q.number}번 문항이 다른 문항과 중복됩니다.")
        stems_seen.add(sn)
        if q.section in (1, 2):
            if q.stem.count("[[BLANK]]") != 1:
                issues.append(f"{q.number}번 빈칸 표시는 [[BLANK]] 정확히 1개여야 합니다.")
            if q.choices:
                issues.append(f"{q.number}번은 빈칸형인데 choices가 있습니다.")
        else:
            if len(q.choices) != 4:
                issues.append(f"{q.number}번 선택지가 4개가 아닙니다.")
            if q.correct_choice not in (1, 2, 3, 4):
                issues.append(f"{q.number}번 정답 선택지 번호가 없습니다.")
            elif len(q.choices) == 4 and norm(q.choices[q.correct_choice - 1]) != tw:
                issues.append(f"{q.number}번 correct choice와 target_word가 일치하지 않습니다.")
        # New Example exact/near reuse check
        for ex in examples:
            if not ex:
                continue
            ratio = SequenceMatcher(None, sn, ex).ratio()
            if sn == ex or ratio >= 0.90:
                issues.append(f"{q.number}번이 Vocabulary List의 New Example과 너무 유사합니다.")
                break
    return issues


# -----------------------------
# DOCX template operations
# -----------------------------

def set_paragraph_text_preserve(p, text: str) -> None:
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def set_cell_text_preserve(cell, text: str) -> None:
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    # remove extra paragraphs while preserving the first paragraph formatting
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def ensure_solid_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = tcPr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tcPr.append(borders)
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = borders.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    borders.append(el)
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "6")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "666666")


def resize_data_rows(table, desired_data_rows: int) -> None:
    current = len(table.rows) - 1
    if current < desired_data_rows:
        prototype = table.rows[-1]._tr
        for _ in range(desired_data_rows - current):
            table._tbl.append(deepcopy(prototype))
    elif current > desired_data_rows:
        for _ in range(current - desired_data_rows):
            tr = table.rows[-1]._tr
            table._tbl.remove(tr)


def portrait_check(doc: Document) -> bool:
    for sec in doc.sections:
        if sec.orientation == WD_ORIENT.LANDSCAPE:
            return False
        # Defensive: width should not exceed height for portrait templates
        if sec.page_width > sec.page_height:
            return False
    return True


def save_list_docx(title: str, unit_title: str, vocab: list[VocabularyItem], path: Path) -> None:
    shutil.copy2(LIST_TEMPLATE, path)
    doc = Document(path)
    set_paragraph_text_preserve(doc.paragraphs[0], title)
    set_paragraph_text_preserve(doc.paragraphs[1], f"{unit_title} - Vocabulary List")
    set_paragraph_text_preserve(doc.paragraphs[2], "※ 예문은 단어 학습을 위한 새 문장이며 시험에는 재사용하지 않습니다. 정확한 동의어·반의어가 없으면 —로 표시합니다.")
    table = doc.tables[0]
    resize_data_rows(table, len(vocab))
    ensure_solid_borders(table)
    for i, v in enumerate(vocab, 1):
        vals = [str(i), v.word, f"{v.pronunciation} {v.pos}", v.meaning_ko, v.syn_ant or "—", v.new_example]
        for j, val in enumerate(vals):
            set_cell_text_preserve(table.cell(i, j), val)
    if not portrait_check(doc):
        raise RuntimeError("Vocabulary List 템플릿의 페이지 방향이 세로가 아닙니다.")
    doc.save(path)


def save_test_docx(title: str, unit_title: str, bundle: TestBundle, path: Path) -> None:
    shutil.copy2(TEST_TEMPLATE, path)
    doc = Document(path)
    set_paragraph_text_preserve(doc.paragraphs[0], title)
    set_paragraph_text_preserve(doc.paragraphs[1], f"{unit_title} - Vocabulary Review Test")
    set_paragraph_text_preserve(doc.paragraphs[2], "학원/반: ____________________    이름: ____________________    점수: ______ / 30")
    set_paragraph_text_preserve(doc.paragraphs[3], "1·2: 첫 글자를 참고하여 알맞은 단어/표현을 쓰시오.   3·4: 가장 알맞은 답을 고르시오.")
    headings = {4: "1. Reading Review (1-12)", 17: "2. Vocabulary Transfer (13-21)", 27: "3. English Definition (22-26)", 33: "4. Vocabulary Relations (27-30)"}
    for idx, text in headings.items():
        set_paragraph_text_preserve(doc.paragraphs[idx], text)
    para_map = list(range(5, 17)) + list(range(18, 27)) + list(range(28, 33)) + list(range(34, 38))
    ordered = sorted(bundle.questions, key=lambda q: q.number)
    for pidx, q in zip(para_map, ordered):
        prompt, _ = render_question(q)
        set_paragraph_text_preserve(doc.paragraphs[pidx], f"{q.number}. {prompt}")
    if not portrait_check(doc):
        raise RuntimeError("Vocabulary Test 템플릿의 페이지 방향이 세로가 아닙니다.")
    doc.save(path)


def save_answer_docx(title: str, unit_title: str, bundle: TestBundle, path: Path) -> None:
    shutil.copy2(ANSWER_TEMPLATE, path)
    doc = Document(path)
    set_paragraph_text_preserve(doc.paragraphs[0], title)
    set_paragraph_text_preserve(doc.paragraphs[1], f"{unit_title} - Answer & Explanation")
    table = doc.tables[0]
    ensure_solid_borders(table)
    ordered = sorted(bundle.questions, key=lambda q: q.number)
    for i, q in enumerate(ordered, 1):
        _, answer = render_question(q)
        vals = [str(i), answer, q.meaning_ko, q.explanation_ko]
        for j, val in enumerate(vals):
            set_cell_text_preserve(table.cell(i, j), val)
    if not portrait_check(doc):
        raise RuntimeError("Answer 템플릿의 페이지 방향이 세로가 아닙니다.")
    doc.save(path)


def convert_docx_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    exe = shutil.which("libreoffice") or shutil.which("soffice")
    if not exe:
        raise RuntimeError("LibreOffice가 설치되어 있지 않아 PDF 변환을 할 수 없습니다.")
    cmd = [exe, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)]
    proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, timeout=120)
    pdf = out_dir / (docx_path.stem + ".pdf")
    if proc.returncode != 0 or not pdf.exists():
        raise RuntimeError(f"PDF 변환 실패: {proc.stderr or proc.stdout}")
    return pdf


def build_export_bundle(title: str, unit_title: str, vocab: list[VocabularyItem], test: TestBundle) -> tuple[dict[str, bytes], list[str]]:
    issues = validate_vocab(vocab) + validate_test(test, vocab)
    if issues:
        return {}, issues
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        paths = {
            "Vocabulary_List.docx": td / "Vocabulary_List.docx",
            "Vocabulary_Test.docx": td / "Vocabulary_Test.docx",
            "Answer_Explanation.docx": td / "Answer_Explanation.docx",
        }
        save_list_docx(title, unit_title, vocab, paths["Vocabulary_List.docx"])
        save_test_docx(title, unit_title, test, paths["Vocabulary_Test.docx"])
        save_answer_docx(title, unit_title, test, paths["Answer_Explanation.docx"])
        files: dict[str, bytes] = {}
        for name, p in paths.items():
            files[name] = p.read_bytes()
        # PDFs
        try:
            for name, p in paths.items():
                pdf = convert_docx_to_pdf(p, td)
                files[pdf.name] = pdf.read_bytes()
        except Exception as e:
            issues.append(str(e))
        # zip
        zbuf = io.BytesIO()
        with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as z:
            for name, data in files.items():
                z.writestr(name, data)
        files["All_Files.zip"] = zbuf.getvalue()
        return files, issues


